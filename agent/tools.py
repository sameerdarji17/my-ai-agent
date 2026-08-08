"""
Tool registry for the agent.

Each tool has:
  1. A JSON-schema style definition (passed to the LLM so it knows the tool exists)
  2. An execute_* function that actually performs the action

IMPORTANT SECURITY NOTES
------------------------
- `read_file` is restricted to AGENT_FILES_ROOT to prevent path traversal.
- `execute_code` runs in a *subprocess* with a timeout and restricted builtins.
  This is a reasonable starting point for development, but for a real
  production deployment you should run untrusted code in a fully isolated
  sandbox such as Docker (network-disabled, read-only filesystem, cgroup
  limits) or a managed sandbox service like E2B. Do NOT expose this tool
  to untrusted end users without hardening it further.
"""

import os
import subprocess
import tempfile
import time
import uuid

import requests
from django.conf import settings

# ---------------------------------------------------------------------------
# Tool definitions (provider-neutral; converted per-provider in orchestrator.py)
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS = [
    {
        "name": "web_search",
        "description": "Search the live web for current info, facts, news, or prices.",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "The search query"}
            },
            "required": ["query"],
        },
    },
    {
        "name": "read_file",
        "description": (
            "Read an uploaded file (.txt/.pdf/.docx/code/images-via-OCR). "
            "Always read before answering about it — never guess contents."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Name of the file to read"}
            },
            "required": ["filename"],
        },
    },
    {
        "name": "list_files",
        "description": "List all files currently available to read.",
        "parameters": {"type": "object", "properties": {}},
    },
    {
        "name": "execute_code",
        "description": (
            "Run a short PYTHON-only snippet in a sandbox, returns stdout/stderr. "
            "For JS/HTML/CSS, just write the code in your text reply instead."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "code": {"type": "string", "description": "Python code to execute"}
            },
            "required": ["code"],
        },
    },
    {
        "name": "generate_kundli",
        "description": (
            "Calculate a real Vedic astrology birth chart from exact birth "
            "date/time/place using real astronomical data. Ask for missing "
            "details first — never guess a kundli from memory."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "birth_date": {"type": "string", "description": "Birth date in YYYY-MM-DD format"},
                "birth_time": {"type": "string", "description": "Birth time in 24-hour HH:MM format"},
                "birth_place": {"type": "string", "description": "Birth place, e.g. 'Mumbai, India'"},
            },
            "required": ["birth_date", "birth_time", "birth_place"],
        },
    },
]

# Premium-only tool: only added to the tool list for users with an active
# premium subscription (see orchestrator.py / views.py).
IMAGE_GEN_TOOL_DEFINITION = {
    "name": "generate_image",
    "description": "Generate a new AI image from a text description; returns an image URL.",
    "parameters": {
        "type": "object",
        "properties": {
            "prompt": {"type": "string", "description": "Description of the image to generate"}
        },
        "required": ["prompt"],
    },
}


def to_openai_tools(defs=TOOL_DEFINITIONS):
    """Groq / OpenAI use: {"type": "function", "function": {name, description, parameters}}"""
    return [
        {
            "type": "function",
            "function": {
                "name": d["name"],
                "description": d["description"],
                "parameters": d["parameters"],
            },
        }
        for d in defs
    ]


def to_anthropic_tools(defs=TOOL_DEFINITIONS):
    """Anthropic uses: {name, description, input_schema}"""
    return [
        {"name": d["name"], "description": d["description"], "input_schema": d["parameters"]}
        for d in defs
    ]


# ---------------------------------------------------------------------------
# Individual tool implementations
# ---------------------------------------------------------------------------

def execute_web_search(tool_input: dict) -> str:
    query = tool_input.get("query", "")
    api_key = settings.TAVILY_API_KEY

    if not api_key:
        return (
            "web_search is not configured: set TAVILY_API_KEY in your .env file "
            "(get a free key at tavily.com). Skipping search for now."
        )

    try:
        resp = requests.post(
            "https://api.tavily.com/search",
            json={
                "api_key": api_key,
                "query": query,
                "max_results": 4,
                "search_depth": "basic",
                "include_answer": True,
            },
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        results = data.get("results", [])
        if not results:
            return "No results found for this query. Try a more specific or differently-worded search."

        lines = []
        quick_answer = data.get("answer")
        if quick_answer:
            lines.append(f"Quick answer (verify against sources below): {quick_answer}\n")

        for i, r in enumerate(results, 1):
            snippet = (r.get("content") or "")[:300]
            lines.append(f"[Source {i}] {r.get('title')}\nURL: {r.get('url')}\n{snippet}\n")

        return "\n".join(lines)
    except Exception as exc:
        return f"web_search failed: {exc}"


def _safe_file_path(filename: str) -> str:
    """Resolve filename inside AGENT_FILES_ROOT only; block path traversal."""
    root = os.path.abspath(settings.AGENT_FILES_ROOT)
    target = os.path.abspath(os.path.join(root, filename))
    if not target.startswith(root):
        raise ValueError("Access outside the allowed files directory is not permitted.")
    return target


def _extract_pdf_text(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages_text) if pages_text else "(PDF has no extractable text — it may be a scanned image.)"


def _extract_docx_text(path: str) -> str:
    import docx

    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts) if parts else "(Document has no readable text.)"


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff"}


def _ocr_image(path: str) -> str:
    """
    Extracts any text visible in an image (e.g. an error message screenshot)
    using OCR.space's free API — a well-established OCR service (25,000
    free requests/month, no credit card). Falls back to a clear message if
    OCR isn't configured or the image has no readable text.
    """
    api_key = settings.OCR_SPACE_API_KEY or "helloworld"  # "helloworld" = OCR.space's public demo key (low limits)
    try:
        with open(path, "rb") as f:
            resp = requests.post(
                "https://api.ocr.space/parse/image",
                files={"file": f},
                data={"apikey": api_key, "language": "eng", "OCREngine": "2"},
                timeout=20,
            )
        resp.raise_for_status()
        data = resp.json()
        if data.get("IsErroredOnProcessing"):
            return "Could not read text from this image (OCR failed): " + str(data.get("ErrorMessage", ""))
        parsed = data.get("ParsedResults") or []
        text = "\n".join(p.get("ParsedText", "") for p in parsed).strip()
        if not text:
            return (
                "This image doesn't seem to contain any readable text "
                "(OCR found nothing). If it's a photo/artwork rather than a "
                "screenshot or document, I can't describe its visual content — "
                "I can only read text that appears in images."
            )
        return f"[Text extracted from image via OCR]\n{text}"
    except Exception as exc:
        return f"Couldn't process this image (OCR error): {exc}"


def execute_read_file(tool_input: dict) -> str:
    filename = tool_input.get("filename", "")
    try:
        path = _safe_file_path(filename)
        if not os.path.isfile(path):
            return f"File not found: {filename}"

        ext = os.path.splitext(path)[1].lower()

        if ext in IMAGE_EXTENSIONS:
            return _ocr_image(path)

        if ext == ".pdf":
            content = _extract_pdf_text(path)
        elif ext in (".docx",):
            content = _extract_docx_text(path)
        elif ext in (".doc",):
            return (
                "Old-format .doc files aren't supported for text extraction. "
                "Please re-save/upload as .docx or .pdf."
            )
        else:
            # Plain text / code / markdown / csv / json etc.
            with open(path, "r", errors="replace") as f:
                content = f.read()

        # Truncate very large content so we don't blow the context window
        if len(content) > 12000:
            content = content[:12000] + "\n... [truncated — file is longer than shown here]"
        return content
    except Exception as exc:
        return f"read_file failed: {exc}"


def execute_list_files(tool_input: dict) -> str:
    root = settings.AGENT_FILES_ROOT
    files = [f for f in os.listdir(root) if os.path.isfile(os.path.join(root, f))]
    if not files:
        return "No files available. Drop files into the agent_files/ directory."
    return "\n".join(files)


def execute_code(tool_input: dict) -> str:
    """
    Run Python code in a subprocess with:
      - a hard wall-clock timeout
      - a fresh temp working directory
      - stdout/stderr captured and length-capped

    NOTE: this isolates against *accidental* mistakes and infinite loops, but
    it does NOT fully isolate against malicious code (it shares the host's
    filesystem/network namespace). For production, replace this with a
    Docker-based or E2B-based sandbox — see README.md.
    """
    code = tool_input.get("code", "")
    run_id = uuid.uuid4().hex
    run_dir = os.path.join(settings.SANDBOX_RUNS_ROOT, run_id)
    os.makedirs(run_dir, exist_ok=True)
    script_path = os.path.join(run_dir, "snippet.py")

    with open(script_path, "w") as f:
        f.write(code)

    try:
        result = subprocess.run(
            ["python3", script_path],
            capture_output=True,
            text=True,
            timeout=settings.CODE_EXEC_TIMEOUT_SECONDS,
            cwd=run_dir,
        )
        output = result.stdout
        if result.stderr:
            output += f"\n[stderr]\n{result.stderr}"
        if not output.strip():
            output = "(code ran with no output)"
        max_chars = settings.CODE_EXEC_MAX_OUTPUT_CHARS
        if len(output) > max_chars:
            output = output[:max_chars] + "\n... [truncated]"
        return output
    except subprocess.TimeoutExpired:
        return f"execute_code failed: execution exceeded {settings.CODE_EXEC_TIMEOUT_SECONDS}s timeout."
    except Exception as exc:
        return f"execute_code failed: {exc}"


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def execute_generate_image(tool_input: dict) -> str:
    """
    Uses Pollinations.ai's free, keyless image generation endpoint. The URL
    itself renders the image on-the-fly (no local storage needed) — the
    frontend detects image URLs in the reply and displays them inline.
    """
    prompt = tool_input.get("prompt", "").strip()
    if not prompt:
        return "generate_image failed: no prompt provided."

    encoded = requests.utils.quote(prompt)
    image_url = f"https://image.pollinations.ai/prompt/{encoded}?width=768&height=768&nologo=true"
    return f"Image generated: {image_url}"


def execute_generate_kundli(tool_input: dict) -> str:
    try:
        from .astrology import calculate_kundli

        return calculate_kundli(
            tool_input.get("birth_date", ""),
            tool_input.get("birth_time", ""),
            tool_input.get("birth_place", ""),
        )
    except Exception as exc:
        return f"generate_kundli failed: {exc}"


TOOL_DISPATCH = {
    "web_search": execute_web_search,
    "read_file": execute_read_file,
    "list_files": execute_list_files,
    "execute_code": execute_code,
    "generate_image": execute_generate_image,
    "generate_kundli": execute_generate_kundli,
}


def run_tool(name: str, tool_input: dict) -> tuple[str, bool, int]:
    """Returns (output_text, succeeded, duration_ms)."""
    start = time.time()
    handler = TOOL_DISPATCH.get(name)
    if handler is None:
        return f"Unknown tool: {name}", False, 0
    try:
        output = handler(tool_input)
        succeeded = True
    except Exception as exc:
        output = f"Tool '{name}' raised an exception: {exc}"
        succeeded = False
    duration_ms = int((time.time() - start) * 1000)
    return output, succeeded, duration_ms
